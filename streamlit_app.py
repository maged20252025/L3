import streamlit as st
import streamlit.components.v1 as components
from docx import Document
import re
import os
import time
import base64
import sqlite3
import uuid
import hashlib

st.set_page_config(page_title="القوانين اليمنية بآخر تعديلاتها حتى عام 2025م", layout="wide")
st.markdown("<h1 style='text-align: center;'>مرحبًا بك في تطبيق القوانين اليمنية بآخر تعديلاتها حتى عام 2025م</h1>", unsafe_allow_html=True)

TRIAL_DURATION = 300 # 5 minutes in seconds
DATABASE_FILE = "user_data.db"
# تم تغيير كلمة مرور المدير هنا
ADMIN_PASSWORD = "maged223562" # تم تغيير كلمة المرور بناءً على طلبك!

def hash_password(password):
    """Hashing a password for simple protection."""
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(stored_hash, provided_password):
    """Verifying a password."""
    return stored_hash == hash_password(provided_password)

def init_db():
    """تهيئة قاعدة البيانات وإنشاء الجداول إذا لم تكن موجودة."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS users (
            user_id TEXT PRIMARY KEY,
            is_activated INTEGER DEFAULT 0,
            trial_start_time REAL,
            last_activity_time REAL,
            activation_code_used TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS activation_codes (
            code TEXT PRIMARY KEY,
            is_used INTEGER DEFAULT 0,
            used_by_user_id TEXT,
            FOREIGN KEY (used_by_user_id) REFERENCES users(user_id)
        )
    ''')
    conn.commit()
    conn.close()

def get_user_id():
    """الحصول على معرف المستخدم الحالي أو إنشائه إذا لم يكن موجودًا."""
    if 'user_id' not in st.session_state:
        st.session_state.user_id = str(uuid.uuid4())
        conn = sqlite3.connect(DATABASE_FILE)
        c = conn.cursor()
        c.execute("INSERT OR IGNORE INTO users (user_id) VALUES (?)", (st.session_state.user_id,))
        conn.commit()
        conn.close()
    return st.session_state.user_id

def update_last_activity(user_id):
    """تحديث وقت آخر نشاط للمستخدم."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("UPDATE users SET last_activity_time = ? WHERE user_id = ?", (time.time(), user_id))
    conn.commit()
    conn.close()

def is_activated(user_id):
    """التحقق مما إذا كان المستخدم مفعلًا من قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("SELECT is_activated FROM users WHERE user_id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    return result[0] == 1 if result else False

def get_trial_start_time(user_id):
    """الحصول على وقت بدء تجربة المستخدم من قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("SELECT trial_start_time FROM users WHERE user_id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    return result[0] if result else None

def set_trial_start_time(user_id):
    """تعيين وقت بدء تجربة المستخدم في قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("INSERT OR IGNORE INTO users (user_id) VALUES (?)", (user_id,))
    c.execute("UPDATE users SET trial_start_time = ?, is_activated = 0 WHERE user_id = ?", (time.time(), user_id))
    conn.commit()
    conn.close()

def activate_app(user_id, code):
    """تفعيل التطبيق للمستخدم باستخدام كود من قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    
    c.execute("SELECT is_used FROM activation_codes WHERE code = ?", (code,))
    code_status = c.fetchone()

    if code_status and code_status[0] == 0: # الكود موجود وغير مستخدم
        try:
            c.execute("UPDATE activation_codes SET is_used = 1, used_by_user_id = ? WHERE code = ?", (user_id, code))
            c.execute("INSERT OR IGNORE INTO users (user_id) VALUES (?)", (user_id,))
            c.execute("UPDATE users SET is_activated = 1, trial_start_time = NULL, activation_code_used = ? WHERE user_id = ?", (code, user_id))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            st.error(f"حدث خطأ أثناء التفعيل: {e}")
            conn.rollback()
            conn.close()
            return False
    else:
        conn.close()
        return False

# ----- وظائف لوحة التحكم (تم نقلها من control_panel_app.py) -----
def generate_activation_code():
    """توليد كود تفعيل فريد."""
    return str(uuid.uuid4()).replace('-', '')[:10].upper()

def save_activation_code(code):
    """حفظ كود التفعيل في قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO activation_codes (code, is_used) VALUES (?, 0)", (code,))
        conn.commit()
        st.success(f"تم توليد الكود: {code} وحفظه بنجاح.")
    except sqlite3.IntegrityError:
        st.warning(f"الكود {code} موجود بالفعل (تجنب التكرار).")
    conn.close()

def get_all_activation_codes():
    """جلب جميع أكواد التفعيل من قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("SELECT code, is_used, used_by_user_id FROM activation_codes ORDER BY is_used ASC, code ASC")
    codes = c.fetchall()
    conn.close()
    return codes

def get_all_users():
    """جلب جميع بيانات المستخدمين من قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("SELECT user_id, is_activated, trial_start_time, last_activity_time, activation_code_used FROM users")
    users = c.fetchall()
    conn.close()
    return users

# لوحة تحكم المدير
def admin_panel():
    st.sidebar.markdown("## ⚙️ لوحة تحكم المدير")
    st.sidebar.markdown("---")

    st.sidebar.header("توليد أكواد التفعيل")
    num_codes = st.sidebar.number_input("عدد الأكواد لتوليدها:", min_value=1, max_value=100, value=1)
    if st.sidebar.button("توليد وحفظ الأكواد"):
        for _ in range(num_codes):
            new_code = generate_activation_code()
            save_activation_code(new_code)
        st.sidebar.success("تم توليد الأكواد المطلوبة!")
        st.rerun() # لإعادة تحميل وعرض الأكواد الجديدة

    st.sidebar.markdown("---")
    st.sidebar.header("عرض الأكواد والمستخدمين")

    # عرض أكواد التفعيل
    st.subheader("قائمة أكواد التفعيل:")
    codes_data = get_all_activation_codes()
    if codes_data:
        with st.expander("عرض/إخفاء أكواد التفعيل", expanded=True):
            for code in codes_data:
                code_str, is_used, used_by = code
                status = "مستخدم" if is_used else "غير مستخدم"
                used_by_text = f"بواسطة: {used_by}" if used_by else "لا يوجد"
                if is_used:
                    st.info(f"**الكود:** `{code_str}`\n\n**الحالة:** {status}\n\n**مستخدم:** {used_by_text}")
                else:
                    st.success(f"**الكود:** `{code_str}`\n\n**الحالة:** {status}")
                st.markdown("---")
    else:
        st.info("لا توجد أكواد تفعيل بعد.")

    # عرض المستخدمين
    st.subheader("قائمة المستخدمين:")
    users_data = get_all_users()
    if users_data:
        with st.expander("عرض/إخفاء المستخدمين", expanded=True):
            for user in users_data:
                user_id, is_activated, trial_start_time, last_activity_time, activation_code_used = user
                status = "مفعل" if is_activated else ("تجريبي" if trial_start_time else "غير مفعل")
                last_activity = time.ctime(last_activity_time) if last_activity_time else "لا يوجد"
                trial_start = time.ctime(trial_start_time) if trial_start_time else "لا يوجد"
                
                if is_activated:
                    st.success(f"**معرف المستخدم:** `{user_id}`\n\n**الحالة:** {status}\n\n**الكود المستخدم:** {activation_code_used if activation_code_used else 'لا يوجد'}\n\n**آخر نشاط:** {last_activity}")
                elif trial_start_time:
                    st.info(f"**معرف المستخدم:** `{user_id}`\n\n**الحالة:** {status}\n\n**بداية التجربة:** {trial_start}\n\n**آخر نشاط:** {last_activity}")
                else:
                    st.warning(f"**معرف المستخدم:** `{user_id}`\n\n**الحالة:** {status}\n\n**آخر نشاط:** {last_activity}")
                st.markdown("---")
    else:
        st.info("لا توجد بيانات مستخدمين بعد.")
    
    st.markdown("---")
    if st.button("تحديث البيانات (لوحة التحكم)"):
        st.rerun()

# ----- وظائف البحث الرئيسية (هي نفسها) -----
def highlight_keywords(text, keywords):
    text = str(text)
    text = text.replace('\xa0', ' ').replace('\u200b', '') 
    
    for kw in keywords:
        text = re.sub(f"({re.escape(kw)})", r"<mark>\1</mark>", text, flags=re.IGNORECASE | re.UNICODE)
    return text

def extract_context(paragraphs, keywords, context_lines=3): 
    paragraphs = [str(p).replace('\xa0', ' ').replace('\u200b', '') for p in paragraphs] 
    
    search_pattern = re.compile('|'.join([re.escape(kw) for kw in keywords]), re.IGNORECASE | re.UNICODE)
    
    matched_indexes = []
    for i, line in enumerate(paragraphs):
        if search_pattern.search(line):
            matched_indexes.append(i)
            
    context_set = set()
    for idx in matched_indexes:
        for i in range(max(0, idx - context_lines), min(len(paragraphs), idx + context_lines + 1)):
            context_set.add(i)
            
    filtered_paragraphs = [paragraphs[i] for i in sorted(context_set) if paragraphs[i].strip()]
    return "\n".join(filtered_paragraphs)

def export_results_to_docx(results, filename="نتائج_البحث.docx"):
    doc = Document()
    doc.add_heading("نتائج البحث", 0)
    for r in results:
        doc.add_heading(f'{r["law"]} - المادة {r["num"]}', level=1)
        doc.add_paragraph(r["context"])
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    return filepath

def run_main_app_logic():
    components.html("""
    <style>
    .scroll-btn {
        position: fixed;
        left: 10px;
        padding: 12px;
        font-size: 24px;
        border-radius: 50%;
        background-color: #c5e1a5;
        color: black;
        cursor: pointer;
        z-index: 9999;
        border: none;
        box-shadow: 1px 1px 5px #888;
    }
    #scroll-top-btn { bottom: 80px; }
    #scroll-bottom-btn { bottom: 20px; }
    </style>
    <button class='scroll-btn' id='scroll-top-btn' onclick='window.scrollTo({top: 0, behavior: "smooth"});'>⬆️</button>
    <button class='scroll-btn' id='scroll-bottom-btn' onclick='window.scrollTo({top: document.body.scrollHeight, behavior: "smooth"});'>⬇️</button>
    """, height=1)

    subfolders = [f.path for f in os.scandir() if f.is_dir() and f.name not in [".git", ".streamlit"]]
    if not subfolders:
        st.warning("📂 لا توجد مجلدات قوانين. يرجى التأكد من وجود مجلدات (مثل 'Laws') تحتوي على ملفات .docx في المستودع.")
        return

    selected_folder = st.selectbox("اختر مجلدًا للبحث فيه:", ["🔍 كل المجلدات"] + subfolders)

    all_files = {}
    if selected_folder == "🔍 كل المجلدات":
        for folder in subfolders:
            files = [f for f in os.listdir(folder) if f.endswith(".docx")]
            all_files[folder] = files
    else:
        files = [f for f in os.listdir(selected_folder) if f.endswith(".docx")]
        all_files[selected_folder] = files

    keywords = st.text_area("الكلمات المفتاحية (افصل بفاصلة)", "")

    if "results" not in st.session_state:
        st.session_state.results = []
    if "search_done" not in st.session_state:
        st.session_state.search_done = False

    if st.button("🔍 بدء البحث") and keywords:
        kw_list = [k.strip() for k in keywords.split(",") if k.strip()]
        results = []

        for folder, files in all_files.items():
            for file in files:
                doc_path = os.path.join(folder, file)
                try:
                    doc = Document(doc_path)
                except Exception as e:
                    st.warning(f"⚠️ تعذر قراءة الملف {file} في المجلد {folder}: {e}. قد يكون الملف تالفًا أو مشفرًا.")
                    continue

                law_name = file.replace(".docx", "")
                
                all_paragraphs_in_doc = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                current_article_paragraphs = []
                last_article_num = "غير معروفة"

                for i, para_text in enumerate(all_paragraphs_in_doc):
                    match = re.match(r"مادة\s*\(?\s*(\d+)\)?", para_text)
                    if match:
                        if current_article_paragraphs:
                            full_article_text = "\n".join(current_article_paragraphs)
                            if any(kw.lower() in full_article_text.lower() for kw in kw_list):
                                context = extract_context(current_article_paragraphs, kw_list, context_lines=3) 
                                results.append({
                                    "law": law_name,
                                    "num": last_article_num,
                                    "text": highlight_keywords(context, kw_list),
                                    "plain": full_article_text,
                                    "context": context,
                                    "keywords": kw_list
                                })
                            current_article_paragraphs = []
                        last_article_num = match.group(1)
                        current_article_paragraphs.append(para_text)
                    else:
                        current_article_paragraphs.append(para_text)
                
                if current_article_paragraphs:
                    full_article_text = "\n".join(current_article_paragraphs)
                    if any(kw.lower() in full_article_text.lower() for kw in kw_list):
                        context = extract_context(current_article_paragraphs, kw_list, context_lines=3) 
                        results.append({
                            "law": law_name,
                            "num": last_article_num,
                            "text": highlight_keywords(context, kw_list),
                            "plain": full_article_text,
                            "context": context,
                            "keywords": kw_list
                        })

        st.session_state.results = results
        st.session_state.search_done = True

    if st.session_state.search_done and st.session_state.results:
        results = st.session_state.results
        unique_laws = sorted(set(r["law"] for r in results))
        st.success(f"تم العثور على {len(results)} نتيجة في {len(unique_laws)} قانون/ملف.")
        
        selected_law = st.selectbox("فلترة حسب القانون", ["الكل"] + unique_laws)
        filtered = results if selected_law == "الكل" else [r for r in results if r["law"] == selected_law]

        for r in filtered:
            st.markdown(f"""
<div style="background-color:#f1f8e9;padding:15px;margin-bottom:15px;border-radius:10px;
            border:1px solid #c5e1a5;direction:rtl;text-align:right; overflow-wrap: break-word;">
    <p style="font-weight:bold;font-size:18px;margin:0">🔷 {r["law"]} - المادة {r["num"]}</p>
    <p style="font-size:17px;line-height:1.8;margin-top:10px">
        {r["text"]}
    </p>
</div>
""", unsafe_allow_html=True)

        if filtered:
            filepath = export_results_to_docx(filtered)
            with open(filepath, "rb") as f:
                st.download_button(
                    label="📥 تحميل النتائج كملف Word",
                    data=f,
                    file_name="نتائج_البحث.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# الوظيفة الرئيسية التي تدير تدفق التطبيق
def main():
    init_db()
    user_id = get_user_id()
    update_last_activity(user_id)

    # إضافة خيار للوصول إلى لوحة المدير في الشريط الجانبي
    st.sidebar.title("خيارات التطبيق")
    app_mode = st.sidebar.radio("اختر الوضع:", ["التطبيق الرئيسي", "لوحة المدير"])

    # تهيئة حالة تسجيل دخول المدير إذا لم تكن موجودة
    if 'admin_logged_in' not in st.session_state:
        st.session_state.admin_logged_in = False

    if app_mode == "لوحة المدير":
        st.sidebar.markdown("---")
        # إذا لم يكن المدير مسجلاً الدخول بعد في هذه الجلسة، اطلب كلمة المرور
        if not st.session_state.admin_logged_in:
            password_input = st.sidebar.text_input("كلمة مرور المدير:", type="password", key="admin_password_input")
            if st.sidebar.button("تسجيل الدخول كمدير"):
                if verify_password(hash_password(ADMIN_PASSWORD), password_input):
                    st.session_state.admin_logged_in = True
                    st.success("تم تسجيل الدخول كمدير بنجاح!")
                    st.rerun() # لإعادة تحميل الصفحة وعرض لوحة المدير مباشرة
                else:
                    st.sidebar.error("كلمة مرور غير صحيحة.")
        
        # إذا كان المدير مسجلاً الدخول (أو بعد تسجيل الدخول بنجاح)
        if st.session_state.admin_logged_in:
            admin_panel()
            if st.sidebar.button("تسجيل الخروج كمدير"):
                st.session_state.admin_logged_in = False
                st.rerun() # لإعادة تحميل الصفحة والعودة لوضع طلب كلمة المرور
    elif app_mode == "التطبيق الرئيسي":
        # إعادة تعيين حالة تسجيل دخول المدير عند التبديل إلى التطبيق الرئيسي
        # st.session_state.admin_logged_in = False # لا حاجة لهذا هنا إذا أردنا الاحتفاظ بـ admin_logged_in عبر تبديل الراديو
        
        if "activated" not in st.session_state:
            st.session_state.activated = is_activated(user_id)

        if not st.session_state.activated:
            st.warning("⚠️ التطبيق غير مفعل. يرجى التفعيل أو استخدام النسخة التجريبية.")
            
            trial_start_time = get_trial_start_time(user_id)
            
            if trial_start_time is None:
                if st.button("🕒 بدء التجربة المجانية", key="start_trial_button"):
                    set_trial_start_time(user_id)
                    st.session_state.trial_start_time = time.time()
                    st.success("🎉 بدأت النسخة التجريبية. لديك 5 دقائق. يرجى تحديث الصفحة أو إعادة تشغيل التطبيق.")
                    st.rerun() 
            else:
                time_elapsed = time.time() - trial_start_time
                if time_elapsed < TRIAL_DURATION:
                    remaining_minutes = int((TRIAL_DURATION - time_elapsed) / 60)
                    st.info(f"✅ النسخة التجريبية نشطة. تبقى لديك حوالي {remaining_minutes} دقيقة.")
                    run_main_app_logic()
                else:
                    st.error("❌ انتهت مدة التجربة المجانية. يرجى التفعيل.")
            
            code = st.text_input("أدخل كود التفعيل هنا", key="activation_code_input")
            if st.button("🔐 تفعيل التطبيق", key="activate_button"):
                if code and activate_app(user_id, code.strip()):
                    st.success("✅ تم التفعيل بنجاح! يرجى تحديث الصفحة أو إعادة تشغيل التطبيق لتطبيق التغييرات.")
                    st.session_state.activated = True
                    st.rerun() 
                else:
                    st.error("❌ كود التفعيل غير صحيح أو تم استخدامه مسبقًا.")
        else:
            run_main_app_logic()

main()
